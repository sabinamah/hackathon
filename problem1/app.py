import os, json, io, csv, re
import fitz
from PIL import Image
from docx import Document as DocxDocument
import streamlit as st
from google import genai
from google.genai import types
from dotenv import load_dotenv

load_dotenv()
api_key_env = os.getenv("GOOGLE_API_KEY", "")

DEPARTMENTS = {
    "IT & Software":      {"icon": "💻", "color": "#3b82f6", "light": "#eff6ff", "border": "#bfdbfe"},
    "Finance":            {"icon": "💰", "color": "#10b981", "light": "#ecfdf5", "border": "#a7f3d0"},
    "Facilities":         {"icon": "🏢", "color": "#f59e0b", "light": "#fffbeb", "border": "#fde68a"},
    "HR & Travel":        {"icon": "✈️",  "color": "#8b5cf6", "light": "#f5f3ff", "border": "#ddd6fe"},
    "Marketing":          {"icon": "🎨", "color": "#ec4899", "light": "#fdf2f8", "border": "#fbcfe8"},
    "Operations":         {"icon": "⚙️",  "color": "#6b7280", "light": "#f9fafb", "border": "#e5e7eb"},
    "Legal & Consulting": {"icon": "⚖️",  "color": "#0891b2", "light": "#ecfeff", "border": "#a5f3fc"},
    "Unknown":            {"icon": "❓", "color": "#94a3b8", "light": "#f8fafc", "border": "#e2e8f0"},
}

SYSTEM_PROMPT = """You are an intelligent invoice processing agent for Globus Group, a large German retail company.

Your job is to read any invoice document in any language and:
1. Extract all key invoice data
2. Categorize the invoice into the correct internal department
3. Assess if it needs urgent attention

Department routing rules:
- "IT & Software"      → software licenses, cloud services, SaaS, hardware, IT subscriptions (Microsoft, AWS, Adobe, Dell, etc.)
- "Finance"            → banking, insurance, tax, financial services
- "Facilities"         → utilities (gas, electricity, water), internet, telephone, building maintenance (Stadtwerke, E.ON, Telekom, etc.)
- "HR & Travel"        → hotels, flights, travel expenses, recruitment, training
- "Marketing"          → advertising, design, print, media, creative services
- "Operations"         → office supplies, logistics, shipping, general supplies
- "Legal & Consulting" → legal fees, consulting, advisory services
- "Unknown"            → cannot determine

Flag as urgent if: total > €10,000 OR due date within 7 days of 2026-06-20 OR overdue.

Respond ONLY with this JSON:
{
  "vendor_name": "string",
  "invoice_number": "string or null",
  "invoice_date": "DD.MM.YYYY or null",
  "due_date": "DD.MM.YYYY or null",
  "currency": "EUR or USD or etc",
  "subtotal": "string or null",
  "vat_amount": "string or null",
  "vat_rate": "string or null",
  "total_amount": "string or null",
  "total_eur_approx": number or null,
  "department": "one of the exact department names above",
  "department_reason": "one sentence explaining WHY this department was chosen",
  "category": "brief description e.g. Cloud Services / Gas Bill / Hotel Stay",
  "confidence": integer 0-100,
  "is_urgent": true or false,
  "urgent_reason": "string or null",
  "language": "German or English or Other",
  "payment_method": "string or null",
  "iban": "string or null",
  "notes": "any important details or null",
  "summary": "2-3 sentences for an accounts payable clerk"
}"""


def extract_content(file_bytes: bytes, filename: str) -> tuple[list[bytes], str]:
    ext = filename.rsplit(".", 1)[-1].lower()
    if ext == "pdf":
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        return [p.get_pixmap(dpi=180).tobytes("png") for p in doc], ""
    elif ext in ("jpg", "jpeg", "png", "webp"):
        img = Image.open(io.BytesIO(file_bytes))
        if img.mode in ("RGBA", "P"):
            img = img.convert("RGB")
        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=92)
        return [buf.getvalue()], ""
    elif ext == "docx":
        doc = DocxDocument(io.BytesIO(file_bytes))
        lines = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    lines.append(row_text)
        full_text = "\n".join(lines)
        # Pre-tag key lines
        tagged = []
        for line in lines:
            l = line.strip()
            if re.search(r"(invoice|rechnung|inv).{0,25}(no|nr|num|number|nummer|#)\s*[:\-=]?\s*\S+", l, re.I) \
               or re.search(r"(no|nr|#)\s*[:\-=]\s*[A-Z0-9\-]{3,}", l, re.I):
                tagged.append(f"[INVOICE NUMBER]: {l}")
            if re.search(r"(invoice|rechnung).{0,20}(date|datum)", l, re.I) \
               or re.search(r"(date|datum|issued|ausgestellt)\s*[:\-=]", l, re.I):
                tagged.append(f"[INVOICE DATE]: {l}")
            if re.search(r"(due|f[äa]llig|payment\s*date|zahlbar|pay\s*by)", l, re.I):
                tagged.append(f"[DUE DATE]: {l}")
            if re.search(r"\d{1,2}[./]\d{1,2}[./]\d{2,4}", l) or re.search(r"\d{4}[.\-/]\d{2}[.\-/]\d{2}", l):
                tagged.append(f"[DATE FOUND]: {l}")
            if re.search(r"(subtotal|netto|net\s*amount)\s*[:\-=]?\s*[\d€$£]", l, re.I):
                tagged.append(f"[SUBTOTAL]: {l}")
            if re.search(r"(vat|mwst|tax)\s*[:\-=]?\s*[\d€$£%]", l, re.I):
                tagged.append(f"[VAT]: {l}")
            if re.search(r"(total|gesamt|grand\s*total)\s*[:\-=]?\s*[\d€$£]", l, re.I):
                tagged.append(f"[TOTAL]: {l}")
        prefix = "=== KEY FIELDS ===\n" + "\n".join(tagged) + "\n\n=== FULL DOCUMENT ===\n" if tagged else ""
        return [], prefix + full_text
    elif ext == "csv":
        raw = file_bytes.decode("utf-8", errors="ignore")
        try:
            reader = csv.DictReader(io.StringIO(raw))
            rows = list(reader)
            if rows:
                lines = ["=== CSV INVOICE DATA ==="]
                lines.append("COLUMNS: " + " | ".join(rows[0].keys()))
                lines.append("")
                for i, row in enumerate(rows):
                    lines.append(f"--- Row {i+1} ---")
                    for k, v in row.items():
                        if v and v.strip():
                            lines.append(f"  {k}: {v.strip()}")
                lines.append("\n=== AUTO-DETECTED FIELDS ===")
                for row in rows:
                    for k, v in row.items():
                        v = v.strip()
                        if not v:
                            continue
                        kl = k.lower()
                        # dates - any format
                        if re.search(r"\d{1,2}[./\-]\d{1,2}[./\-]\d{2,4}", v) or re.search(r"\d{4}[./\-]\d{2}[./\-]\d{2}", v):
                            lines.append(f"  [DATE] {k} = {v}")
                        # invoice number - any column with number/nr/no/id
                        if re.search(r"(invoice|rechnung|inv|bill|nr|no\b|number|nummer|id)", kl):
                            lines.append(f"  [INVOICE_NUM] {k} = {v}")
                        # due date
                        if re.search(r"(due|f.llig|zahlbar|payment.*date|datum)", kl):
                            lines.append(f"  [DUE_DATE] {k} = {v}")
                        # amounts
                        if re.search(r"(total|gesamt|amount|betrag|summe|subtotal|netto|brutto|vat|mwst|steuer)", kl):
                            lines.append(f"  [AMOUNT] {k} = {v}")
                        # iban
                        if re.search(r"iban|bank|konto", kl):
                            lines.append(f"  [BANK] {k} = {v}")
                return [], "\n".join(lines)
        except Exception:
            pass
        return [], raw
    elif ext == "txt":
        raw = file_bytes.decode("utf-8", errors="ignore")
        hints = []
        for line in raw.splitlines():
            l = line.strip()
            if not l:
                continue
            # Invoice number - broad match
            if re.search(r"(invoice|rechnung|inv|bill).{0,20}(no|nr|num|number|nummer|#)\s*[:\-=]?\s*\S+", l, re.I) \
               or re.search(r"(no|nr|number|nummer)\s*[:\-=]\s*\S+", l, re.I):
                hints.append(f"[INVOICE NUMBER LINE]: {l}")
            # Invoice date - broad match
            if re.search(r"(invoice|rechnung|inv).{0,20}(date|datum|dat)", l, re.I) \
               or re.search(r"(date|datum|ausgestellt|erstellt)\s*[:\-=]", l, re.I):
                hints.append(f"[INVOICE DATE LINE]: {l}")
            # Due date
            if re.search(r"(due|f[äa]llig|zahlbar|payment\s*date|zu zahlen|payment\s*term)", l, re.I):
                hints.append(f"[DUE DATE LINE]: {l}")
            # Any line containing a date pattern
            if re.search(r"\d{1,2}[./]\d{1,2}[./]\d{2,4}", l) or re.search(r"\d{4}[./\-]\d{2}[./\-]\d{2}", l):
                hints.append(f"[DATE FOUND]: {l}")
            # Total/amount
            if re.search(r"(total|gesamt|betrag|amount|summe|netto|brutto|vat|mwst)\s*[:\-=]?\s*[\d€$£]", l, re.I):
                hints.append(f"[AMOUNT LINE]: {l}")
        header = "=== KEY LINES DETECTED ===\n" + "\n".join(hints) + "\n\n=== FULL TEXT ===\n" if hints else "=== FULL TEXT ===\n"
        return [], header + raw
    return [], ""


def process_invoice(api_key: str, file_bytes: bytes, filename: str) -> dict:
    images, text = extract_content(file_bytes, filename)
    client = genai.Client(api_key=api_key)
    ext = filename.rsplit(".", 1)[-1].lower()
    prompt = f"Invoice filename: {filename}\nFile type: {ext.upper()}\n"
    if text:
        prompt += f"\n--- DOCUMENT CONTENT START ---\n{text}\n--- DOCUMENT CONTENT END ---\n"
        if ext == "csv":
            prompt += "\nIMPORTANT: If this CSV has multiple rows, it is an invoice manifest/list — use the FIRST data row as the primary invoice. Extract vendor_name, currency, vat_rate, total_amount from the columns directly. If invoice_number or dates are not in the CSV columns, set them to null. Put a summary of all rows in the 'notes' field. Look for [INVOICE_NUM], [DATE], [DUE_DATE], [AMOUNT] tagged lines for any available values.\n"
        elif ext == "txt":
            prompt += "\nIMPORTANT: Lines tagged [INVOICE NUMBER LINE], [INVOICE DATE LINE], [DUE DATE LINE], [DATE FOUND], [AMOUNT LINE] contain the key data — extract values directly from those lines. For invoice_number: look at [INVOICE NUMBER LINE] and extract the actual number after the colon/equals. For invoice_date: extract the date from [INVOICE DATE LINE] or [DATE FOUND]. Never return null if a tagged line contains the data.\n"
    if ext == "docx":
        prompt += "\nIMPORTANT: Lines tagged [INVOICE NUMBER], [INVOICE DATE], [DUE DATE], [DATE FOUND], [SUBTOTAL], [VAT], [TOTAL] contain the key data. Extract values directly from those lines. Never return null if a tagged line contains the data.\n"
    prompt += "\nExtract and categorize this invoice according to the JSON schema."
    parts = [types.Part.from_text(text=prompt)]
    for img in images[:4]:
        mime = "image/png" if img[:4] == b"\x89PNG" else "image/jpeg"
        parts.append(types.Part.from_bytes(data=img, mime_type=mime))
    response = client.models.generate_content(
        model="gemini-2.5-flash",
        contents=parts,
        config=types.GenerateContentConfig(
            system_instruction=SYSTEM_PROMPT,
            response_mime_type="application/json",
        ),
    )
    raw = response.text.strip().lstrip("```json").lstrip("```").rstrip("```").strip()
    result = json.loads(raw)
    if isinstance(result, list):
        result = result[0] if result else {}
    return result


# ─── Page ─────────────────────────────────────────────────────────────────────
st.set_page_config(page_title="InvoiceAI — Globus Group", page_icon="🧾", layout="wide", initial_sidebar_state="collapsed")

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&display=swap');
*, html, body { font-family: 'Inter', sans-serif !important; }
.main { background: #f0f4f9 !important; }
.block-container { padding: 2rem 3rem !important; max-width: 100% !important; }
#MainMenu, footer, header { visibility: hidden; }

.topbar {
  background: linear-gradient(120deg, #14532d 0%, #15803d 100%);
  padding: 26px 40px; display: flex; align-items: center; gap: 18px;
  border-radius: 16px; margin-bottom: 28px;
}
.topbar-title { color: white; font-size: 1.7rem; font-weight: 800; margin: 0; }
.topbar-sub   { color: rgba(255,255,255,0.75); font-size: 0.88rem; margin: 3px 0 0; }
.topbar-pill  {
  display: inline-block; background: rgba(255,255,255,0.15);
  border: 1px solid rgba(255,255,255,0.3); color: rgba(255,255,255,0.9);
  padding: 3px 13px; border-radius: 999px; font-size: 0.77rem; margin-top: 7px;
}
.panel {
  background: white; border-radius: 16px; border: 1px solid #e2e8f0;
  box-shadow: 0 2px 10px rgba(0,0,0,0.06); padding: 28px;
}
.panel-label {
  font-size: 0.7rem; font-weight: 700; letter-spacing: 0.1em;
  text-transform: uppercase; color: #94a3b8; margin-bottom: 10px;
}
.stButton > button {
  background: linear-gradient(135deg, #14532d, #16a34a) !important;
  color: white !important; border: none !important;
  border-radius: 10px !important; font-weight: 600 !important;
  font-size: 0.95rem !important; padding: 0.65rem !important;
  width: 100% !important; box-shadow: 0 2px 8px rgba(20,83,45,0.3) !important;
}
.stButton > button:disabled { opacity: 0.35 !important; }
.stTextInput > div > div > input {
  border-radius: 9px !important; border: 1.5px solid #dde3ee !important;
  font-size: 0.9rem !important; padding: 10px 14px !important; background: #f8faff !important;
}
[data-testid="stFileUploaderDropzone"] {
  background: #f8fff8 !important; border: 2px dashed #86efac !important; border-radius: 12px !important;
}

/* Department routing card - BIG and clear */
.routing-card {
  border-radius: 16px; padding: 24px 28px; margin-bottom: 20px;
  border: 2px solid;
}
.routing-label {
  font-size: 0.72rem; font-weight: 700; letter-spacing: 0.1em;
  text-transform: uppercase; margin-bottom: 10px; opacity: 0.7;
}
.routing-arrow { font-size: 1rem; margin-right: 6px; }
.routing-dept  { font-size: 1.6rem; font-weight: 800; margin: 0; }
.routing-cat   { font-size: 0.9rem; margin: 4px 0 8px; opacity: 0.75; }
.routing-why   {
  font-size: 0.82rem; background: rgba(255,255,255,0.6);
  border-radius: 8px; padding: 8px 12px; margin-top: 10px; line-height: 1.5;
}

.conf-bar-bg { background: #e5e7eb; border-radius: 999px; height: 10px; margin: 4px 0 6px; }
.urgent-badge {
  display: inline-flex; align-items: center; gap: 6px;
  background: #fef2f2; border: 1px solid #fecaca;
  color: #b91c1c; padding: 7px 16px; border-radius: 999px;
  font-size: 0.83rem; font-weight: 700; margin-bottom: 16px;
}
.detail-grid {
  display: grid; grid-template-columns: 1fr 1fr;
  border: 1px solid #f1f5f9; border-radius: 12px; overflow: hidden; margin-bottom: 18px;
}
.dc { padding: 11px 16px; border-bottom: 1px solid #f1f5f9; }
.dc:nth-child(odd)  { border-right: 1px solid #f1f5f9; background: #fafbff; }
.dc:nth-child(even) { background: white; }
.dc:nth-last-child(-n+2) { border-bottom: none; }
.dc-l { font-size: 0.68rem; font-weight: 700; color: #94a3b8; text-transform: uppercase; letter-spacing: 0.05em; }
.dc-v { font-size: 0.9rem; font-weight: 600; color: #1e293b; margin-top: 2px; }
.summary-box {
  background: #f0fdf4; border-left: 4px solid #16a34a;
  border-radius: 0 10px 10px 0; padding: 13px 17px;
  font-size: 0.88rem; color: #166534; line-height: 1.6; margin-bottom: 18px;
}
.dept-row {
  display: flex; align-items: center; gap: 10px;
  padding: 8px 0; border-bottom: 1px solid #f8f9fa; font-size: 0.84rem;
}
.dept-row:last-child { border-bottom: none; }
.placeholder {
  background: white; border: 2px dashed #bbf7d0; border-radius: 16px;
  padding: 64px 40px; text-align: center;
}
[data-testid="stDownloadButton"] > button {
  background: white !important; color: #166534 !important;
  border: 1.5px solid #86efac !important; border-radius: 10px !important;
  font-weight: 600 !important; font-size: 0.88rem !important;
}
</style>
""", unsafe_allow_html=True)

# ── Header ────────────────────────────────────────────────────────────────────
st.markdown("""
<div class="topbar">
  <div style="font-size:2.4rem">🧾</div>
  <div>
    <p class="topbar-title">InvoiceAI</p>
    <p class="topbar-sub">Automated Invoice Processing &amp; Department Routing · Powered by Google Gemini AI</p>
    <span class="topbar-pill">Problem 1 — Globus Group, St. Wendel</span>
  </div>
</div>
""", unsafe_allow_html=True)

left, right = st.columns([0.9, 1.1], gap="large")

# ── LEFT ──────────────────────────────────────────────────────────────────────
with left:
    st.markdown('<div class="panel">', unsafe_allow_html=True)

    st.markdown('<div class="panel-label">🔑 API Configuration</div>', unsafe_allow_html=True)
    api_key = st.text_input("key", type="password", value=api_key_env,
                            placeholder="AIza...", label_visibility="collapsed")
    st.markdown('<p style="font-size:0.75rem;color:#94a3b8;margin:-4px 0 18px 2px">Free key at <strong>aistudio.google.com/apikey</strong></p>', unsafe_allow_html=True)

    st.markdown('<div class="panel-label">📎 Upload Invoice</div>', unsafe_allow_html=True)
    uploaded = st.file_uploader("inv", type=["pdf", "png", "jpg", "jpeg", "docx", "csv", "txt"],
                                label_visibility="collapsed")
    if uploaded:
        st.markdown(f'<div style="background:#f0fdf4;border:1px solid #86efac;border-radius:8px;padding:9px 13px;font-size:0.84rem;color:#166534;margin-top:6px">✓ <strong>{uploaded.name}</strong> · {uploaded.size/1024:.1f} KB</div>', unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)
    go = st.button("⚡  Process Invoice", disabled=not uploaded)

    st.markdown('<hr style="border:none;border-top:1px solid #f1f5f9;margin:22px 0">', unsafe_allow_html=True)
    st.markdown('<div class="panel-label">🏢 Department Routing Map</div>', unsafe_allow_html=True)
    for dept, cfg in DEPARTMENTS.items():
        if dept == "Unknown":
            continue
        st.markdown(f'<div class="dept-row"><span style="font-size:1.1rem">{cfg["icon"]}</span><span style="font-weight:600;color:#1e293b">{dept}</span></div>', unsafe_allow_html=True)
    st.markdown('<p style="font-size:0.75rem;color:#94a3b8;margin:14px 0 0">Supports PDF · PNG · JPG · DOCX · CSV · TXT · DE &amp; EN</p>', unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

# ── RIGHT ─────────────────────────────────────────────────────────────────────
with right:
    if go and uploaded:
        if not api_key:
            st.error("Please enter your Google API key.")
        else:
            with st.spinner("Reading and processing invoice…"):
                try:
                    r = process_invoice(api_key, uploaded.read(), uploaded.name)
                    st.session_state["r1"] = r
                    st.session_state["r1fn"] = uploaded.name
                except Exception as e:
                    st.error(f"Processing failed: {e}")

    if "r1" in st.session_state:
        r = st.session_state["r1"]
        dept  = r.get("department", "Unknown")
        dcfg  = DEPARTMENTS.get(dept, DEPARTMENTS["Unknown"])
        conf  = r.get("confidence", 0)
        urgent = r.get("is_urgent", False)

        # ── BIG department routing card ──
        st.markdown(f"""
        <div class="routing-card" style="background:{dcfg['light']};border-color:{dcfg['border']}">
          <div class="routing-label" style="color:{dcfg['color']}">📨 This invoice belongs to</div>
          <p class="routing-dept" style="color:{dcfg['color']}">{dcfg['icon']} &nbsp;{dept} Department</p>
          <p class="routing-cat" style="color:{dcfg['color']}">{r.get('category','')}</p>
          <div class="routing-why">💬 <em>{r.get('department_reason','')}</em></div>
        </div>
        """, unsafe_allow_html=True)

        # Confidence bar
        bar_color = "#16a34a" if conf >= 75 else "#f59e0b" if conf >= 50 else "#ef4444"
        st.markdown(f'<div style="font-size:0.7rem;color:#94a3b8;font-weight:700;letter-spacing:0.05em;text-transform:uppercase">Routing Confidence</div>', unsafe_allow_html=True)
        st.markdown(f'<div class="conf-bar-bg"><div style="background:{bar_color};width:{conf}%;height:10px;border-radius:999px"></div></div>', unsafe_allow_html=True)
        st.markdown(f'<div style="font-size:0.8rem;color:{bar_color};font-weight:700;margin:-4px 0 16px">{conf}% confidence</div>', unsafe_allow_html=True)

        if urgent:
            st.markdown(f'<div class="urgent-badge">🚨 URGENT — {r.get("urgent_reason","Requires immediate attention")}</div>', unsafe_allow_html=True)

        if r.get("summary"):
            st.markdown(f'<div class="summary-box">{r["summary"]}</div>', unsafe_allow_html=True)

        # Invoice details
        def dc(label, value):
            return f'<div class="dc"><div class="dc-l">{label}</div><div class="dc-v">{value or "—"}</div></div>'

        total = r.get("total_amount") or "—"
        total_eur = r.get("total_eur_approx")
        total_str = f'{total}{"  (~€{:,.0f})".format(total_eur) if total_eur and r.get("currency","EUR") != "EUR" else ""}'

        st.markdown(f"""
        <div class="detail-grid">
          {dc("Vendor", r.get("vendor_name"))}
          {dc("Invoice No.", r.get("invoice_number"))}
          {dc("Invoice Date", r.get("invoice_date"))}
          {dc("Due Date", r.get("due_date"))}
          {dc("Currency", r.get("currency"))}
          {dc("VAT Rate", r.get("vat_rate"))}
          {dc("Subtotal", r.get("subtotal"))}
          {dc("VAT Amount", r.get("vat_amount"))}
          <div class="dc" style="grid-column:span 2;background:#f0fdf4">
            <div class="dc-l">Total Amount</div>
            <div class="dc-v" style="font-size:1.1rem;color:#166534">{total_str}</div>
          </div>
          {dc("Language", r.get("language"))}
          {dc("Payment", r.get("payment_method"))}
          {f'{dc("IBAN", r.get("iban"))}' if r.get("iban") else ""}
          {f'{dc("Notes", r.get("notes"))}' if r.get("notes") else ""}
        </div>
        """, unsafe_allow_html=True)

        st.download_button(
            label="📥  Download Invoice Report (JSON)",
            data=json.dumps(r, indent=2, ensure_ascii=False),
            file_name=f"invoice_{st.session_state.get('r1fn','report')}.json",
            mime="application/json",
            use_container_width=True,
        )

    else:
        st.markdown("""
        <div class="placeholder">
          <div style="font-size:3.5rem;margin-bottom:14px">🧾</div>
          <div style="font-size:1.05rem;font-weight:700;color:#374151;margin-bottom:8px">No invoice uploaded yet</div>
          <div style="font-size:0.875rem;color:#94a3b8;max-width:300px;margin:0 auto;line-height:1.6">
            Upload any supplier invoice — the AI reads it, extracts all data,
            and tells you exactly which department should handle it.
          </div>
          <div style="margin-top:24px;display:flex;gap:10px;flex-wrap:wrap;justify-content:center">
            <span style="background:#f0fdf4;color:#166534;border:1px solid #86efac;padding:5px 14px;border-radius:999px;font-size:0.8rem;font-weight:600">⚡ Instant routing</span>
            <span style="background:#f0fdf4;color:#166534;border:1px solid #86efac;padding:5px 14px;border-radius:999px;font-size:0.8rem;font-weight:600">🌍 DE &amp; EN</span>
            <span style="background:#f0fdf4;color:#166534;border:1px solid #86efac;padding:5px 14px;border-radius:999px;font-size:0.8rem;font-weight:600">📊 VAT extracted</span>
            <span style="background:#f0fdf4;color:#166534;border:1px solid #86efac;padding:5px 14px;border-radius:999px;font-size:0.8rem;font-weight:600">🚨 Urgent flagging</span>
          </div>
        </div>
        """, unsafe_allow_html=True)
